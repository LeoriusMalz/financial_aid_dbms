import os

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert

class DownloadFile:
    def __init__(self, tmp_dir: str = "data/download/tmp"):
        self.tmp_dir = tmp_dir

    def __set_font_settings(self, paragraph, font_name='Times New Roman', font_size=14, alignment=None):
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)

        if alignment:
            if alignment == 'center':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif alignment == 'right':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif alignment == 'left':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif alignment == 'justify':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def __fill_application_template_styled(self, output_path, data, template_path="data/app_template.docx"):
        doc = Document(template_path)

        new_doc = Document()
        in_loop = False

        for paragraph in doc.paragraphs:
            text = paragraph.text

            if '{% for category in categories %}' in text:
                in_loop = True
                continue

            if '{% endfor %}' in text:
                in_loop = False

                heading_para = new_doc.add_paragraph("Прошу оказать мне материальную помощь по причине:")
                self.__set_font_settings(heading_para)

                for i, category in enumerate(data['categories'], 1):
                    reason_para = new_doc.add_paragraph(f"{i}. {category}")
                    self.__set_font_settings(reason_para)
                continue

            if in_loop:
                continue

            text = text.replace('{{group}}', data['group'])
            text = text.replace('{{fullname}}', data['fullname'])
            text = text.replace('{{phone}}', data['phone'])
            text = text.replace('{{d}}', data['d'])
            text = text.replace('{{m}}', data['m'])
            text = text.replace('{{Y}}', data['Y'])

            if '{{category}}' in text:
                text = text.replace('- {{category}}', '')

            if text.strip():
                new_para = new_doc.add_paragraph(text)

                if 'Директору школы' in text:
                    self.__set_font_settings(new_para, alignment='right')
                elif 'Заявление' in text:
                    self.__set_font_settings(new_para, alignment='center')
                else:
                    self.__set_font_settings(new_para)

        new_doc.save(output_path)
        return output_path

    def word(self, name, data):
        return self.__fill_application_template_styled(os.path.join(self.tmp_dir, f"{name}.docx"), data)

    def pdf(self, name, data):
        convert(self.word(name, data), os.path.join(self.tmp_dir, f"{name}.pdf"))

# df = DownloadFile()
# data = {
#         'group': 'ФРКТ-123',
#         'fullname': 'Мальцев Лев Александрович',
#         'phone': '+7 (999) 123-45-67',
#         'categories': [
#             'Сложное материальное положение в семье',
#             'Необходимость оплаты дополнительного образования',
#             'Расходы на медицинское обслуживание'
#         ],
#         'd': '1',
#         'm': 'декабря',
#         'Y': '2024'
#     }
#
# df.word("test2", data)
# df.pdf("Мальцев", data)

# if __name__ == "__main__":
#     data = {
#         'group': 'ФРКТ-123',
#         'fullname': 'Мальцев Лев Александрович',
#         'phone': '+7 (999) 123-45-67',
#         'categories': [
#             'Сложное материальное положение в семье',
#             'Необходимость оплаты дополнительного образования',
#             'Расходы на медицинское обслуживание'
#         ],
#         'd': '1',
#         'm': 'декабря',
#         'Y': '2024'
#     }

    # fill_application_template_styled(
    #     template_path='app_template.docx',
    #     output_path='../../Мальцев Лев Александрович_заполненный.docx',
    #     data=data
    # )